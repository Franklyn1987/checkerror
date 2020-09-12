#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Date    : 2020-04-07 20:22:19
# @Author  : 崔立波 (baiyuexingchen@gmail.com)
# @Link    : http://blog.sina.com.cn/dejavu1
# @Version : 1



import requests
import time
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE
from docx.enum.section import WD_ORIENT
import correct_text as c
###############################################
#以下开始
def doc():
    print("checkerror('文件路径')，输出至D:\\")
def checkerror(path):
	strslice=path.split("\\",-1)
	for x in enumerate(strslice):
		if x[1].find(".do")>0:
			wenjianming=x[1]
	mk=wenjianming.find(".docx")
	wenjianming=wenjianming[0:mk]
	document=Document(path)
	for x in document.paragraphs:
		# print(x.text)
		if x.text=="      ":
			continue
		if x.text=="    ":
			continue
		if x.text=="\r\n":
			continue
		if x.text.strip()=="":
			continue
		try:
			c.txt_correction(x.text)
		except Exception as e:
			print("此处无法识别")
		

	try:

		document=Document(path)
		# print(1)
	except Exception as e:
		document=Document()

	for paragraph in document.paragraphs:
	        for run in paragraph.runs:
	            if "%" in run.text:
	                run.font.name = u'Times New Roman'
	                print("修改%")
	            if "." in run.text:
	                run.font.name = u'Times New Roman'
	                print("修改.")
	            if "," in run.text:
	            	run.text=run.text.replace(",","，")
	            	run.font.name = u'仿宋_GB2312'
	            	print("修改，")
	            if ")" in run.text:
	            	run.text=run.text.replace(")","）")
	            	# run.font.name = u'仿宋_GB2312'
	            	print("修改)")
	            if "(" in run.text:
	            	run.text=run.text.replace("(","（")
	            	# run.font.name = u'仿宋_GB2312'
	            	print("修改(")
	            if "[" in run.text:
	                run.font.name = u'仿宋_GB2312'
	                print("修改[")
	            if "]" in run.text:
	                run.font.name = u'仿宋_GB2312'
	                print("修改]")
	pa="D:/"+wenjianming+time.strftime('%Y-%m-%d-%H-%M-%S', time.localtime(time.time()))+".docx"
	print(pa)
	document.save("D:/"+wenjianming+time.strftime('%Y-%m-%d-%H-%M-%S', time.localtime(time.time()))+".docx")